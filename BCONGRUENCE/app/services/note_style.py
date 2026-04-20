"""
Note Style Processing Service

Handles extraction of text from uploaded clinical notes (PDF, DOC, DOCX, TXT)
and provides utilities for style matching in note generation.

MVP Implementation - focuses on reliable text extraction and simple style matching.
"""

import logging
import base64
import io
from typing import Optional, Dict, Any
from datetime import datetime

logger = logging.getLogger(__name__)

def extract_text_from_file(file_content: str, file_type: str) -> str:
    """
    Extract text from uploaded file.
    
    Args:
        file_content: Base64 encoded file content
        file_type: File extension (pdf, doc, docx, txt)
    
    Returns:
        Extracted text content
    
    Raises:
        ValueError: If file type is unsupported
        Exception: If extraction fails
    """
    try:
        # Decode base64 content
        file_bytes = base64.b64decode(file_content)
        
        if file_type.lower() == 'pdf':
            return extract_text_from_pdf(file_bytes)
        elif file_type.lower() == 'docx':
            return extract_text_from_docx(file_bytes)
        elif file_type.lower() == 'txt':
            return file_bytes.decode('utf-8')
        elif file_type.lower() == 'doc':
            # For MVP, we'll skip legacy .doc support
            raise ValueError("Legacy .doc files not supported in MVP. Please convert to .docx or .pdf")
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
            
    except Exception as e:
        logger.error(f"Failed to extract text from {file_type}: {e}")
        raise

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyPDF2."""
    try:
        import PyPDF2
        
        pdf_file = io.BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        extracted_text = text.strip()
        
        if not extracted_text:
            raise ValueError("No text could be extracted from PDF. The file may be image-based or corrupted.")
        
        logger.info(f"Extracted {len(extracted_text)} characters from PDF")
        return extracted_text
        
    except ImportError:
        raise ValueError("PyPDF2 not installed. Please install with: pip install PyPDF2")
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"Failed to extract text from PDF: {e}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        
        doc_file = io.BytesIO(file_bytes)
        doc = Document(doc_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        extracted_text = text.strip()
        
        if not extracted_text:
            raise ValueError("No text could be extracted from DOCX file.")
        
        logger.info(f"Extracted {len(extracted_text)} characters from DOCX")
        return extracted_text
        
    except ImportError:
        raise ValueError("python-docx not installed. Please install with: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {e}")

def get_preview_text(full_text: str, max_chars: int = 200) -> str:
    """
    Get preview of note text for display purposes.
    
    Args:
        full_text: Full note text
        max_chars: Maximum characters for preview
    
    Returns:
        Preview text with ellipsis if truncated
    """
    if not full_text or not full_text.strip():
        return "No preview available"
    
    # Clean up the text
    clean_text = full_text.strip()
    
    if len(clean_text) <= max_chars:
        return clean_text
    
    # Try to break at a sentence or line
    preview = clean_text[:max_chars]
    last_period = preview.rfind('.')
    last_newline = preview.rfind('\n')
    
    # Find the best break point
    break_point = max(last_period, last_newline)
    if break_point > max_chars * 0.7:  # If we found a good break point
        return clean_text[:break_point + 1] + "..."
    else:
        return preview + "..."

def validate_note_content(note_text: str) -> Dict[str, Any]:
    """
    Validate extracted note content for quality and usability.
    
    Args:
        note_text: Extracted note text
    
    Returns:
        Dictionary with validation results
    """
    validation = {
        "is_valid": False,
        "char_count": len(note_text),
        "word_count": len(note_text.split()),
        "has_structure": False,
        "warnings": [],
        "errors": []
    }
    
    # Basic length validation
    if len(note_text.strip()) < 50:
        validation["errors"].append("Note content too short (minimum 50 characters)")
        return validation
    
    if len(note_text) > 50000:  # 50KB limit
        validation["warnings"].append("Note content very long - may be truncated in processing")
    
    # Check for basic clinical note structure
    note_lower = note_text.lower()
    structure_indicators = [
        'subjective:', 'objective:', 'assessment:', 'plan:',  # SOAP
        'chief complaint:', 'history of present illness:', 'hpi:',  # Standard
        'mental status:', 'diagnosis:', 'treatment:', 'medication:',  # Clinical
        'patient reports:', 'patient states:', 'client reports:'  # Common phrases
    ]
    
    found_indicators = [indicator for indicator in structure_indicators if indicator in note_lower]
    if found_indicators:
        validation["has_structure"] = True
        validation["structure_indicators"] = found_indicators
    else:
        validation["warnings"].append("No clear clinical note structure detected")
    
    # Check for potential issues
    if note_text.count('\n') < 3:
        validation["warnings"].append("Note appears to be single paragraph - may lack structure")
    
    # Overall validation
    validation["is_valid"] = len(validation["errors"]) == 0
    
    return validation

def analyze_note_style(note_text: str) -> Dict[str, Any]:
    """
    Analyze the style characteristics of a clinical note.
    
    Args:
        note_text: Clinical note text
    
    Returns:
        Dictionary with style analysis
    """
    analysis = {
        "format_type": "unknown",
        "sections": [],
        "tone": "unknown",
        "detail_level": "unknown",
        "avg_sentence_length": 0,
        "section_count": 0
    }
    
    lines = note_text.split('\n')
    
    # Detect sections (lines that end with : or are all caps)
    sections = []
    for line in lines:
        line_stripped = line.strip()
        if line_stripped and (line_stripped.endswith(':') or line_stripped.isupper()):
            sections.append(line_stripped)
    
    analysis["sections"] = sections
    analysis["section_count"] = len(sections)
    
    # Detect format type
    section_text = ' '.join(sections).lower()
    if all(keyword in section_text for keyword in ['subjective', 'objective', 'assessment', 'plan']):
        analysis["format_type"] = "SOAP"
    elif 'chief complaint' in section_text or 'hpi' in section_text:
        analysis["format_type"] = "Traditional"
    elif 'mental status' in section_text:
        analysis["format_type"] = "Psychiatric"
    
    # Analyze tone and detail level
    sentences = [s.strip() for s in note_text.replace('\n', ' ').split('.') if s.strip()]
    if sentences:
        avg_length = sum(len(s.split()) for s in sentences) / len(sentences)
        analysis["avg_sentence_length"] = round(avg_length, 1)
        
        if avg_length < 8:
            analysis["tone"] = "concise"
            analysis["detail_level"] = "brief"
        elif avg_length < 15:
            analysis["tone"] = "moderate"
            analysis["detail_level"] = "standard"
        else:
            analysis["tone"] = "detailed"
            analysis["detail_level"] = "comprehensive"
    
    return analysis

def prepare_style_context(note_text: str, max_chars: int = 2000) -> str:
    """
    Prepare note text for use as style context in LLM prompts.
    
    Args:
        note_text: Full note text
        max_chars: Maximum characters to include
    
    Returns:
        Truncated note text suitable for prompt context
    """
    if len(note_text) <= max_chars:
        return note_text
    
    # Try to include complete sections rather than cutting mid-section
    lines = note_text.split('\n')
    context = ""
    
    for line in lines:
        if len(context + line + '\n') > max_chars:
            break
        context += line + '\n'
    
    return context.strip()